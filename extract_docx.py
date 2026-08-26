import docx, os

output_path = r'c:/Users/KIIT0001/Desktop/BillItUp/all_docx_content.txt'
docx_files = [f for f in os.listdir('.') if f.endswith('.docx')]

with open(output_path, 'w', encoding='utf-8', errors='ignore') as f:
    for docx_file in docx_files:
        f.write(f"--- CONTENT OF {docx_file} ---\n")
        try:
            doc = docx.Document(docx_file)
            f.write("PARAGRAPHS:\n")
            for p in doc.paragraphs:
                f.write(p.text + '\n')
            
            f.write("\nTABLES:\n")
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    f.write(" | ".join(row_text) + "\n")
                f.write("-" * 20 + "\n")
        except Exception as e:
            f.write(f"Error reading {docx_file}: {e}\n")
        f.write("\n\n")

print(f'Extraction of {len(docx_files)} files complete.')
